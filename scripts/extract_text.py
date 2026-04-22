#!/usr/bin/env python3
"""
extract_text.py — Извлечение текста из документов.

Поддерживает: PDF (с текстовым слоем и OCR), DOCX, TXT, изображения (через tesseract).
Возвращает текст + метаданные (метод извлечения, уверенность, количество страниц).

Использование:
    python3 extract_text.py <путь_к_файлу> [--output-dir <папка>]

Выход:
    Печатает JSON с полями: text, method, confidence, pages, warnings
    Если --output-dir указан — также сохраняет .txt рядом.
"""

import sys
import json
import os
import subprocess
from pathlib import Path

def extract_pdf_text(filepath: str) -> dict:
    """Извлечение текста из PDF. Сначала программно, потом OCR."""
    try:
        import fitz  # pymupdf
    except ImportError:
        return {"text": "", "method": "none", "confidence": "low",
                "pages": 0, "warnings": ["pymupdf не установлен"]}

    doc = fitz.open(filepath)
    pages = len(doc)
    text_parts = []
    has_text = False

    for page in doc:
        page_text = page.get_text()
        if page_text.strip():
            has_text = True
        text_parts.append(page_text)

    doc.close()

    if has_text:
        full_text = "\n\n---\n\n".join(text_parts)
        # Оценка качества: если средняя длина страницы > 100 символов — хорошо
        avg_len = len(full_text) / max(pages, 1)
        confidence = "high" if avg_len > 200 else "medium" if avg_len > 50 else "low"
        return {
            "text": full_text,
            "method": "pdf-text",
            "confidence": confidence,
            "pages": pages,
            "warnings": []
        }

    # Нет текстового слоя — пробуем OCR
    return extract_pdf_ocr(filepath, pages)


def extract_pdf_ocr(filepath: str, pages: int) -> dict:
    """OCR через tesseract для PDF без текста."""
    try:
        result = subprocess.run(
            ["ocrmypdf", "--force-ocr", "-l", "rus+eng", "--sidecar", "/dev/stdout",
             filepath, "/dev/null"],
            capture_output=True, text=True, timeout=120
        )
        if result.stdout.strip():
            return {
                "text": result.stdout,
                "method": "ocr",
                "confidence": "medium",
                "pages": pages,
                "warnings": ["Текст извлечён через OCR, возможны ошибки"]
            }
    except (subprocess.TimeoutExpired, FileNotFoundError):
        pass

    # Fallback: tesseract напрямую (для одностраничных)
    try:
        result = subprocess.run(
            ["tesseract", filepath, "stdout", "-l", "rus+eng"],
            capture_output=True, text=True, timeout=60
        )
        if result.stdout.strip():
            return {
                "text": result.stdout,
                "method": "ocr",
                "confidence": "low",
                "pages": pages,
                "warnings": ["Tesseract fallback, качество может быть низким"]
            }
    except (subprocess.TimeoutExpired, FileNotFoundError):
        pass

    return {
        "text": "",
        "method": "none",
        "confidence": "low",
        "pages": pages,
        "warnings": ["Не удалось извлечь текст. Требуется LLM fallback (Read tool)."]
    }


def extract_docx_text(filepath: str) -> dict:
    """Извлечение текста из DOCX."""
    try:
        from docx import Document
    except ImportError:
        return {"text": "", "method": "none", "confidence": "low",
                "pages": 0, "warnings": ["python-docx не установлен"]}

    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)

    # Таблицы
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            text += "\n\n" + "\n".join(rows)

    return {
        "text": text,
        "method": "docx-parse",
        "confidence": "high",
        "pages": max(1, len(paragraphs) // 30),  # грубая оценка
        "warnings": []
    }


def extract_text_file(filepath: str) -> dict:
    """Чтение текстового файла."""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            text = f.read()
    except UnicodeDecodeError:
        try:
            with open(filepath, "r", encoding="cp1251") as f:
                text = f.read()
        except Exception:
            return {"text": "", "method": "none", "confidence": "low",
                    "pages": 0, "warnings": ["Не удалось прочитать файл"]}

    return {
        "text": text,
        "method": "text-read",
        "confidence": "high",
        "pages": 1,
        "warnings": []
    }


def extract_image_ocr(filepath: str) -> dict:
    """OCR для изображений через tesseract."""
    try:
        result = subprocess.run(
            ["tesseract", filepath, "stdout", "-l", "rus+eng"],
            capture_output=True, text=True, timeout=60
        )
        if result.stdout.strip():
            confidence = "medium" if len(result.stdout) > 100 else "low"
            return {
                "text": result.stdout,
                "method": "ocr",
                "confidence": confidence,
                "pages": 1,
                "warnings": ["Текст из изображения через tesseract"]
            }
    except (subprocess.TimeoutExpired, FileNotFoundError):
        pass

    return {
        "text": "",
        "method": "none",
        "confidence": "low",
        "pages": 1,
        "warnings": ["Не удалось распознать. Требуется LLM fallback (Read tool для изображения)."]
    }


def extract(filepath: str) -> dict:
    """Главная функция: определяет тип файла и вызывает нужный экстрактор."""
    ext = Path(filepath).suffix.lower()

    if ext == ".pdf":
        return extract_pdf_text(filepath)
    elif ext in (".docx",):
        return extract_docx_text(filepath)
    elif ext in (".txt", ".md", ".csv", ".html", ".htm", ".xml", ".json", ".yaml", ".yml"):
        return extract_text_file(filepath)
    elif ext in (".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp"):
        return extract_image_ocr(filepath)
    elif ext in (".doc", ".rtf", ".odt"):
        return {
            "text": "",
            "method": "none",
            "confidence": "low",
            "pages": 0,
            "warnings": [f"Формат {ext} не поддерживается программно. Используй Read tool."]
        }
    else:
        return {
            "text": "",
            "method": "none",
            "confidence": "low",
            "pages": 0,
            "warnings": [f"Неизвестный формат: {ext}"]
        }


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Использование: python3 extract_text.py <файл> [--output-dir <папка>]")
        sys.exit(1)

    filepath = sys.argv[1]
    output_dir = None

    if "--output-dir" in sys.argv:
        idx = sys.argv.index("--output-dir")
        if idx + 1 < len(sys.argv):
            output_dir = sys.argv[idx + 1]

    if not os.path.exists(filepath):
        print(json.dumps({"error": f"Файл не найден: {filepath}"}, ensure_ascii=False))
        sys.exit(1)

    result = extract(filepath)

    if output_dir and result["text"]:
        os.makedirs(output_dir, exist_ok=True)
        txt_path = os.path.join(output_dir, Path(filepath).stem + ".txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(result["text"])
        result["saved_to"] = txt_path

    print(json.dumps(result, ensure_ascii=False, indent=2))
